from copy import deepcopy
from enum import unique
from io import BytesIO
from os import read
from pydoc import doc
import string
from docx import Document
import csv
import io
from docxtpl import DocxTemplate
from docxcompose.composer import Composer
from docx.enum.text import WD_BREAK

from numpy import full

def generate_one_file(csv_file, template_file, added_rows, new_header):
    rows, bad_rows = parse_csv(csv_file)
    if (new_header):
        print(new_header)
        rows[0] = new_header
    if added_rows:
        for i in range(len(added_rows)):
            print(added_rows[i])
            rows.append(added_rows[i])
    path = replace_tokens(rows[0], rows[1], template_file)
    return path


def replace_tokens(tokens, row, template):
    context = dict(zip(tokens, row))
    print(context)
    doc = DocxTemplate(template)
    doc.render(context)
    output_path = row[0] + ".docx"
    doc.save(output_path)
    return output_path


def generate_files_docx(csv_file, template_file, added_rows, new_header):
    rows, bad_rows = parse_csv(csv_file)
    if (new_header):
        print(new_header)
        rows[0] = new_header
    if added_rows:
        for i in range(len(added_rows)):
            print(added_rows[i])
            rows.append(added_rows[i])
    all_filled_documents, paths = replace_tokens_with_values_in_rows(rows, template_file)
    return all_filled_documents, paths


def merge_documents(documents_paths, result_document_path):
    master = Document(documents_paths[0])
    composer = Composer(master)
    for path in documents_paths[1:]:
        doc = Document(path)
        para = master.add_paragraph()
        run = para.add_run()
        run.add_break(WD_BREAK.PAGE)
        composer.append(doc)
    composer.save(result_document_path)


def replace_tokens_with_values_in_rows(rows, template):
    filled_templates = []
    tokens = rows[0]
    paths = []

    print(rows)
    for row in rows[1:]:
        context = dict(zip(tokens, row))
        print(context)
        doc = DocxTemplate(template)
        doc.render(context)
        output_path = row[0] + ".docx"
        paths.append(output_path)
        doc.save(output_path)
        filled_templates.append(Document(output_path))
    return filled_templates, paths


def parse_csv(file, delimiter=';'):
    text = file.read().decode('utf-8-sig', errors='replace')
    reader = csv.reader(io.StringIO(text), delimiter=delimiter)

    rows = []
    bad_rows = []

    header = next(reader)
    expected_count_columns = len(header)
    rows.append(header)
    for line_num, row in enumerate(reader, start=2):
        if len(row) != expected_count_columns:
            bad_rows.append(row)
        else:
            rows.append(row)

    return rows, bad_rows


def csv_to_text(rows, bad_rows):
    buffer = io.StringIO()
    writer = csv.writer(buffer, delimiter=';')

    buffer.write("=== VALID ROWS ===\n")
    for r in rows:
        writer.writerow(r)

    buffer.write("\n=== BAD ROWS ===\n")
    for r in bad_rows:
        writer.writerow(r)

    text = buffer.getvalue()
    buffer.close()

    return text # Для отладки