from io import BytesIO
import json
from django.shortcuts import render
from django.http import HttpResponse, JsonResponse
from docx import Document
from .utils import generate_files_docx, generate_one_file, merge_documents, parse_csv

def generate(request):
    if request.method == "POST":
        template_file = request.FILES.get('template_file')
        csv_file = request.FILES.get('csv_file')

        if not csv_file or not template_file:
            return HttpResponse("Файлы не загружены", status=400)

        added_rows_json = request.POST.get('added_rows')
        added_rows = json.loads(added_rows_json)
        new_header_json = request.POST.get('added_header')
        new_header = json.loads(new_header_json)

        output_format = request.POST.get('output_format')
        output_type = request.POST.get('output_type')

        if output_type == 'file':
            all_filled_documents, paths = generate_files_docx(csv_file, template_file, added_rows, new_header)
            path_output = "merged.docx"
            merge_documents(paths, path_output)
        else:
            ...

        result = Document(path_output)
        output = BytesIO()
        result.save(output)
        output.seek(0)

        filename = "result.docx"

        response = HttpResponse(output.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response

    return render(request, "generator/index.html")


def check_csv(request):
    if request.method == "POST" and request.FILES.get("csv_file"):
        csv_file = request.FILES.get('csv_file')

        rows, bad_rows = parse_csv(csv_file)

        return JsonResponse({
            "header": rows[0],
            "bad_rows": bad_rows,
            "good_rows": rows[1:]
        })

    return JsonResponse({"error": "no file"}, status=400)


def preview_docx(request):
    if request.method == "POST":
        template_file = request.FILES.get('template_file')
        csv_file = request.FILES.get('csv_file')

        if not csv_file or not template_file:
            return HttpResponse("Файлы не загружены", status=400)

        added_rows_json = request.POST.get('added_rows')
        added_rows = json.loads(added_rows_json)
        new_header_json = request.POST.get('added_header')
        new_header = json.loads(new_header_json)

        path = generate_one_file(csv_file, template_file, added_rows, new_header)
        doc = Document(path)

        html = "".join(f"<p>{p.text}</p>" for p in doc.paragraphs)
        print(html)

        return HttpResponse(html, content_type="text/html")